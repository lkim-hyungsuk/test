from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
for section in doc.sections:
    section.page_width = Inches(13)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

# Title
title = doc.add_heading('Full Ranked Matrix — Apartment Options for Residency Fit (OSU Med)', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(14)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x36, 0x64)

doc.add_paragraph()

# Table data
headers = [
    "Rank", "Apartment", "Tracy Quote / Known Monthly", "Commute to OSU Med",
    "Best Pros", "Biggest Cons", "Residency-Fit Verdict"
]

rows = [
    [
        "1",
        "The Edge",
        "Studio $1,490–$1,590\n1-bed $1,830\n(Garage +$125/mo)",
        "1.2 mi / ~5 min",
        "• 2025 build\n• Assigned/attached garage\n• Coworking, clubhouse\n• 24-hr fitness, yoga room\n• Pool, courtyard/firepit\n• Close to Aldi/shopping",
        "• Fee stack is real\n• Train-noise question not fully resolved\n• Fewer reviews (new building)",
        "BEST OVERALL PICK\nClosest + newest + easy garage life."
    ],
    [
        "2",
        "Steel House",
        "1-bed $1,695\n+ own wifi",
        "1.6 mi / ~7 min",
        "• Quartz, stainless, in-unit W/D\n• 9'/10' ceilings, walk-in closets\n• Pool, fitness, Starbucks coffee bar\n• Lounge, fire pit, grills",
        "• Not as new as Edge/Fifth\n• Valet-trash concern\n• Pedestrian/crossing concern on Chambers",
        "BEST VALUE PICK\nSpend less without dropping quality."
    ],
    [
        "3",
        "The Broadview",
        "1-bed / 1.5 bath\n$1,925 base\n(fee sheet TBD)",
        "2.1 mi / ~9 min",
        "• 858 sq ft, 1.5 baths, balcony\n• W/D, oversized kitchen island\n• Premium appliances, 10' ceilings\n• Boutique feel, EV charging\n• Assigned garages",
        "• Fee sheet still missing\n• Utilities not included\n• Farthest of the active four",
        'BEST "ACTUALLY LIVING THERE"\nPrioritize space & finish quality over shortest commute.'
    ],
    [
        "4",
        "Fifth x Northwest",
        "1-bed $1,878\n+ unclear utility cap",
        "1.8 mi / ~7 min",
        "• Structured parking, EV charging\n• Coworking + conference room\n• Sauna, sky deck, pool\n• Coffee bar, fiber internet\n• Highest walkability of the group",
        "• Expensive once fees pile up\n• Utility cap still unclear\n• Thin walls & ungated garage (reviews)",
        "LUXURY/AMENITY HEAVY\nEdge still beats it for residency practicality."
    ],
    [
        "5",
        "Westmont at the Lane",
        "1-bed $1,741\n+ own wifi",
        "3.1 mi",
        "• Quartz, stainless, in-unit W/D\n• Clubhouse/amenity deck\n• Walkable Upper Arlington area\n• Very quiet evenings (resident review)",
        "• Clearly farther from OSU\n• Weaker transit score\n• Less convenient for early/late hours",
        "GOOD FALLBACK\nCalmer vibes, tolerable longer drive."
    ],
    [
        "6",
        "Apts at the Yard: Morrison",
        "1-bed $1,764\n+ own wifi",
        "Grandview Yard area",
        "• Free public garage/surface lot\n• Modern interiors\n• Residents praise location, staff, maintenance",
        "• Official floor plan page didn't load well\n• Near train tracks\n• Multi-building setup is less straightforward",
        "POSSIBLE SLEEPER\nMore uncertainty than top options."
    ],
    [
        "7",
        "Tribeca",
        "1-bed $1,612\n+ possible utilities question",
        "Near Grandview / near tracks",
        "• In-unit W/D, controlled parking garage\n• Pool, lounge, fitness, bike storage\n• Responsive maintenance (reviews)\n• Mostly quiet except trains",
        "• Built 2013 (Tracy dislikes)\n• Train-track concern is real\n• Weaker walkability",
        "BUDGET-FRIENDLY BUT COMPROMISED\nFine if price wins, not if newer + calmer matters."
    ],
    [
        "8",
        "Fairfax",
        "Layout E: $1,984 or $1,525\n+ uncapped communal utilities\n+ own wifi",
        "Kenny Rd",
        "• Newer building\n• $1,525 unit is tempting outlier\n• Staff praised in reviews\n• Strong walk score (Zillow)",
        "• Uncapped communal utilities\n• Final monthly cost is unpredictable",
        "ONLY IF CHEAP UNIT IS STILL AVAILABLE\nAnd utilities turn out manageable."
    ],
    [
        "9",
        "Langham",
        "1-bed $2,004\n+ uncapped communal utilities",
        "Chambers Rd / near OSU",
        "• Smart-home features, coworking\n• Theater, golf simulators\n• Hydrotherapy pool, lounge\n• Smart thermostats, controlled access\n• Very close to OSU Health System",
        "• Pricey\n• Uncapped communal utilities = huge red flag\n• Not predictable enough financially",
        "TEMPTING BUT RISKY\nDo not move up without hard written utility numbers."
    ],
    [
        "10",
        "Essex",
        "1-bed A2 $2,029\n+ uncapped communal utilities\n+ own wifi",
        "Chambers Rd",
        "• W/D, quartz, coffee bar\n• Coworking, Peloton bikes\n• Courtyard, garage\n• Access to Langham amenities\n• Move-in credit marketing",
        "• Same utility-risk as Langham\n• Trash included but most utilities are not\n• Hard to justify vs. Edge/Steel/Broadview",
        "LOWEST PRIORITY\nHardest to justify in this lineup."
    ],
]

# Color scheme
RANK_COLORS = {
    "1": RGBColor(0xC6, 0xEF, 0xCE),   # green
    "2": RGBColor(0xC6, 0xEF, 0xCE),   # green
    "3": RGBColor(0xC6, 0xEF, 0xCE),   # green
    "4": RGBColor(0xFF, 0xEB, 0x9C),   # yellow
    "5": RGBColor(0xFF, 0xEB, 0x9C),   # yellow
    "6": RGBColor(0xFF, 0xEB, 0x9C),   # yellow
    "7": RGBColor(0xFF, 0xEB, 0x9C),   # yellow
    "8": RGBColor(0xFF, 0xC7, 0xCE),   # red/pink
    "9": RGBColor(0xFF, 0xC7, 0xCE),
    "10": RGBColor(0xFF, 0xC7, 0xCE),
}

HEADER_COLOR = RGBColor(0x1F, 0x36, 0x64)
HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_background(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(color)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '4472C4')
        tcBorders.append(border)
    tcPr.append(tcBorders)

# Column widths (in inches), total ~12 inches
col_widths = [0.45, 1.1, 1.35, 1.0, 2.5, 2.3, 2.3]

table = doc.add_table(rows=1 + len(rows), cols=len(headers))
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for i, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = Inches(width)

# Header row
hdr_row = table.rows[0]
for i, h in enumerate(headers):
    cell = hdr_row.cells[i]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_background(cell, HEADER_COLOR)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = HEADER_TEXT_COLOR

# Data rows
for r_idx, row_data in enumerate(rows):
    row = table.rows[r_idx + 1]
    rank = row_data[0]
    bg = RANK_COLORS.get(rank, RGBColor(0xFF, 0xFF, 0xFF))

    for c_idx, cell_text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_background(cell, bg)
        set_cell_borders(cell)

        p = cell.paragraphs[0]
        is_verdict = (c_idx == 6)
        is_rank = (c_idx == 0)
        is_name = (c_idx == 1)

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (is_rank or is_name) else WD_ALIGN_PARAGRAPH.LEFT

        lines = cell_text.split('\n')
        for l_idx, line in enumerate(lines):
            if l_idx == 0:
                run = p.add_run(line)
            else:
                run = p.add_run('\n' + line)
            run.font.size = Pt(8)
            if is_verdict and l_idx == 0:
                run.bold = True
                run.font.size = Pt(8.5)
            if is_rank:
                run.bold = True
                run.font.size = Pt(11)
            if is_name:
                run.bold = True
                run.font.size = Pt(9)

# Legend
doc.add_paragraph()
legend_para = doc.add_paragraph()
legend_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_colored_legend(para, color: RGBColor, label: str):
    run = para.add_run("  " + label + "  ")
    run.font.size = Pt(9)
    run.font.bold = True
    # Highlight via shading requires a table; use font color as fallback
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

legend_title = doc.add_paragraph()
t = legend_title.add_run("Color Key: ")
t.bold = True
t.font.size = Pt(9)

# Create a small legend table
legend_table = doc.add_table(rows=1, cols=3)
legend_table.style = 'Table Grid'
legend_items = [
    (RGBColor(0xC6, 0xEF, 0xCE), "Ranks 1–3: Top Picks (Green)"),
    (RGBColor(0xFF, 0xEB, 0x9C), "Ranks 4–7: Consider Carefully (Yellow)"),
    (RGBColor(0xFF, 0xC7, 0xCE), "Ranks 8–10: Low Priority / Risk (Red)"),
]
for i, (color, label) in enumerate(legend_items):
    cell = legend_table.rows[0].cells[i]
    set_cell_background(cell, color)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(8.5)
    run.bold = True

# Footer note
doc.add_paragraph()
note = doc.add_paragraph(
    "Note: Ranking is based on residency fit — shortest/easiest drive to OSU Med, "
    "parking convenience, predictable monthly cost, building quality, and likelihood of decent sleep. "
    "Drive times are based on Tracy's notes."
)
note.runs[0].font.size = Pt(8)
note.runs[0].italic = True

output_path = "/home/user/test/Apartment_Ranked_Matrix.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
